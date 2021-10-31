import pandas as pd
import os
from collections import defaultdict
import argparse
import docx
from docx.shared import Pt
from tabulate import tabulate
import collections
import eval_utils as utils

looking_at_partial_matches = []
def calc_worker_stats(df):
    '''
    :return: calculated annotation statistics
    '''
    df['len_annotations'] = df['answers'].apply(lambda x: len(x))
    df['avg_num_annotations'] = df.groupby('worker_id')['len_annotations'].transform('mean')


def print_stats(pairs_performance, workerious, df):
    workers = df.groupby(['worker_id', 'hit_count','avg_num_annotations']).size().reset_index()
    workers.drop([0],axis=1, inplace=True)
    workers['precision'] = workers['worker_id'].apply(lambda x: workerious[x][0])
    workers['recall'] = workers['worker_id'].apply(lambda x: workerious[x][1])
    workers['F1'] = workers['worker_id'].apply(lambda x: workerious[x][2])
    print("Workers performance on batch:")
    print(tabulate(workers, headers=workers.columns, tablefmt='orgtbl'))
    print("Avg performance on hits: ")
    print(tabulate(pairs_performance.items(), headers=['hit id', 'avg iou']))


def calc_edges_iou(pred, gold_keys, eval_edges, soft_match):
    '''

    :param pred: workers predictions
    :param gold_keys: expert annotations
    :param eval_edges: do evaluation at edge level
    :param soft_match: add soft match in addition to exact match on alignments/edges
    :return:
    '''
    pairs2performance = defaultdict(list) #overall performance at the HIT level
    worker2feedback = defaultdict(dict) #collecting feedbacks for workers
    worker2iou = defaultdict() #worker's P,R, and F1

    for wrk_num, wrk in enumerate(pred['worker_id'].unique()):
        wrk_df = pred[pred.worker_id == wrk]
        total_alignment_predictions = 0
        total_correct_aligments = 0
        total_gold_predictions = 0  # since each worker might do a different number of alignments
        for i, row in wrk_df.iterrows():

            #if row['key'] != 'DUC2006~!~D0643~!~21~!~DUC2006~!~D0643~!~3': continue
            if row.key not in gold_keys: continue
            curr_gold_keys = gold_keys[row.key]
            if not curr_gold_keys['unique_edges'] and not row['answers']:
                continue
            worker2feedback[wrk][row.key] = {"prec_err": [], "recall_err": [], "correct_aligns": [], "feedback": '{}'}

            correct_aligments, alignment_predictions, gold_predictions, instance_feedback, \
            instance_pairs2performance, partial_matches = utils.calc_alignment_iou(row['answers'], curr_gold_keys, eval_edges, soft_match)

            looking_at_partial_matches.extend(partial_matches)
            total_correct_aligments += correct_aligments
            total_alignment_predictions += alignment_predictions
            total_gold_predictions += gold_predictions

            if (instance_feedback['prec_err'] or instance_feedback['recall_err']) or row['feedback'] != '{}':
                worker2feedback[wrk][row.key] = instance_feedback
                worker2feedback[wrk][row.key]["feedback"] = row['feedback']
            else:
                if wrk == 'A175DRQDVKWQA7':
                    worker2feedback[wrk][row.key] = instance_feedback
                    worker2feedback[wrk][row.key]["feedback"] = row['feedback']
                else:
                    # we dont keep feedbacks for correct alignments
                    del worker2feedback[wrk][row.key]

            pairs2performance[row['key']].append(instance_pairs2performance)

        if total_alignment_predictions == 0:
            prec = 0
        else:
            prec = total_correct_aligments / total_alignment_predictions

        recall = total_correct_aligments / total_gold_predictions

        if (prec + recall) == 0:
            f1 = 0
        else:
            f1 = 2 * (prec * recall) / (prec + recall)
        worker2iou[wrk] = (round(prec, 2), round(recall, 2), round(f1, 2))

    return pairs2performance, worker2feedback, worker2iou


def prep_alignment_fordoc(align, end=False):
    align_str = ""
    align_str += "A: " + str(align['sent1']) + "\n"
    if end:
        align_str += "B: " + str(align['sent2'])
    else:
        align_str += "B: " + str(align['sent2']) + "\n"

    return align_str

def get_bolded_sents(html_sent, paragraph):
    for token in html_sent.split(" "):
        if "strong>" in token:
            run = paragraph.add_run()
            tok = token.replace("<strong>", "")
            tok = tok.replace("</strong>", "")
            run.text = tok + " "
            run.font.bold = True
            continue
        run = paragraph.add_run()
        run.text = token + " "
    return paragraph

def print_worker_eval_docx(worker_feedback, worker2iou, output_dir, gold_sents):
    for wrk, feedback in worker_feedback.items():
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        doc.add_heading("WorkerId: " + wrk, 0)
        doc.add_heading("Main Feedback - Overall Score: " + str(worker2iou[wrk][2]), 1)  # F1 score
        doc.add_paragraph("Thank you for the hard work.")
        doc.add_heading("Errors - Missed or Incorrect Alignments\n", 1)
        counter = 1
        ordered_feedbacks = collections.OrderedDict(sorted(feedback.items()))

        for key, alignments in ordered_feedbacks.items(): #need to sort by keys

            doc.add_paragraph("------------ " + str(counter) + " ------------")
            counter += 1
            sents = gold_sents[key]
            pA = doc.add_paragraph()
            pB = doc.add_paragraph()
            pA = get_bolded_sents("Sentence A: "+sents[0], pA)
            pB = get_bolded_sents("Sentence B: "+sents[1], pB)
            if alignments['feedback'] != '{}':
                pC = doc.add_paragraph("Feedback from worker: " + alignments['feedback'])
                pC.style = doc.styles['Normal']
            pA.style = doc.styles['Normal']
            pB.style = doc.styles['Normal']


            table = doc.add_table(rows=7, cols=1)
            table.style = "LightShading-Accent1"
            table.rows[0].cells[0].text = "- Main Errors - "
            table.rows[1].cells[0].text = 'Incorrect Alignments'
            if not alignments['prec_err']:
                table.rows[2].cells[0].text = "None\n"
            else:
                cell_str = ""
                for j, item in enumerate(alignments['prec_err']):
                    if j == len(alignments['prec_err']) - 1:
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item)
                    else:
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item) + "\n"
                table.rows[2].cells[0].text = cell_str

            table.rows[3].cells[0].text = 'Missed Alignments'
            if not alignments['recall_err']:
                table.rows[4].cells[0].text = "None\n"
            else:
                cell_str = ""
                for j, item in enumerate(alignments['recall_err']):
                    if j == (len(alignments['recall_err']) - 1):
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item)
                    else:
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item) + "\n"
                table.rows[4].cells[0].text = cell_str

            table.rows[5].cells[0].text = 'Correct Alignments'
            if not alignments['correct_aligns']:
                table.rows[6].cells[0].text = "None\n"
            else:
                cell_str = ""
                for j, item in enumerate(alignments['correct_aligns']):
                    if j == len(alignments['correct_aligns']) - 1:
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item)
                    else:
                        cell_str += "(" + str(j + 1) + ") \n" + prep_alignment_fordoc(item) + "\n"
                table.rows[6].cells[0].text = cell_str

            doc.add_heading("Explanation:", 1)
            doc.add_paragraph("\n\n")

        if wrk == 'A175DRQDVKWQA7':
            wrk = 'gold'
        output_path = os.path.join(output_dir, wrk)
        doc.save(output_path + ".docx")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--pred_annots', type=str, required=True,
                        help='csv file containing qa alignment annotations by workers')
    parser.add_argument('--gold', type=str, required=True,
                        help='csv file containing gold qa alignments')
    parser.add_argument('--eval_dir', type=str, required=True,
                        help='path to directory to save workers evaluations')
    parser.add_argument('--soft_match', action="store_true", required=False, default=False,
                        help='evaluate soft alignment matches in addition to exact match')
    parser.add_argument('--edge_level', action="store_true", required=False, default=False,
                        help='evaluate at the edge level in a bi-partite graph, rather than alignment level')
    parser.add_argument('--save', action="store_true", required=False, default=False,
                        help='path to directory to save workers evaluations')

    args = parser.parse_args()
    pred = pd.read_csv(args.pred_annots)
    gold = pd.read_csv(args.gold)

    pred['answers'] = pred['answers'].apply(lambda x: eval(x)) #csv files should be after 'post_qaalign' script
    gold['answers'] = gold['answers'].apply(lambda x: eval(x))

    calc_worker_stats(pred) #avg num annotations / batch

    gold_keys = utils.get_gold_key_edges(gold, args.edge_level) #create gold dictionary for batch
    gold_sents = dict(zip(gold['key'], list(zip(gold['text_1_html'], gold['text_2_html'])))) #for .docx files

    pairs2performance, worker2feedback, worker2iou = calc_edges_iou(pred, gold_keys, args.edge_level, args.soft_match)

    pairs2performance = {k: sum(v) / len(v) for k, v in pairs2performance.items()}

    if args.save:
        print("Saving workers feedbacks")
        print_worker_eval_docx(worker2feedback, worker2iou, args.eval_dir, gold_sents)

    print("Batch Stats")
    print_stats(pairs2performance, worker2iou, pred)

    gold_sents = dict(zip(gold['key'], list(zip(gold['text_1'], gold['text_2']))))

    pairs_performance_df = pd.DataFrame(pairs2performance.items(), columns=['key', 'avg_iou'])
    pairs_performance_df['text_1'] = pairs_performance_df['key'].apply(lambda x: gold_sents[x][0])
    pairs_performance_df['text_2'] = pairs_performance_df['key'].apply(lambda x: gold_sents[x][1])
    pairs_performance_df.to_csv(os.path.join(args.eval_dir, 'avg_hit_performance.csv'), index=False)

    pd.DataFrame(looking_at_partial_matches).to_csv(args.eval_dir+"/partial_matches.csv", index=False)